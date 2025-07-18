import streamlit as st
from docx import Document
from PyPDF2 import PdfReader
import tempfile
import os
from openai import OpenAI
from base64 import b64encode

# إعداد الصفحة
st.set_page_config(page_title="منصة إعداد العروض - متوازي", layout="centered")

# تحويل صورة الشعار إلى base64 لعرضها في الزاوية
def get_base64_logo(image_path):
    with open(image_path, "rb") as image_file:
        return b64encode(image_file.read()).decode()

logo_base64 = get_base64_logo("logo_corner.png")

# تحديد اللغة
lang = st.selectbox("Language / اللغة", ["العربية", "English"])

# نصوص الواجهة حسب اللغة
TEXTS = {
    "العربية": {
        "title": "منصة إعداد العروض - متوازي",
        "instruction": "قم برفع كراسة الشروط وسيتم توليد عرض فني احترافي",
        "upload_label": "ارفع كراسة الشروط (PDF)",
        "project_name": "اسم المشروع",
        "client_name": "اسم الجهة",
        "gov_logo": "شعار الجهة الحكومية (اختياري)",
        "generate_button": "توليد العرض الفني",
        "loading": "جارٍ قراءة الكراسة وتحليلها...",
        "download": "تحميل العرض الفني (Word)",
        "success": "تم توليد العرض الفني بنجاح!",
        "error": "يرجى تعبئة جميع الحقول المطلوبة.",
        "doc_heading": "العرض الفني للمشروع"
    },
    "English": {
        "title": "Proposal Generation Platform - Mutawazi",
        "instruction": "Upload the RFP and a professional technical proposal will be generated",
        "upload_label": "Upload RFP file (PDF)",
        "project_name": "Project Name",
        "client_name": "Client Name",
        "gov_logo": "Client Logo (optional)",
        "generate_button": "Generate Proposal",
        "loading": "Reading and analyzing the document...",
        "download": "Download Technical Proposal (Word)",
        "success": "Proposal generated successfully!",
        "error": "Please fill in all required fields.",
        "doc_heading": "Technical Proposal for Project"
    }
}

# إدراج الشعار والثيم
st.markdown(
    f"""
    <style>
    .stApp {{
        background-color: #e6f4ea;
    }}
    h1 {{
        color: #004d26;
    }}
    .stTextInput > div > div > input {{
        background-color: #f7fcf9;
        border: 1px solid #aad4bc;
    }}
    .stDownloadButton button {{
        background-color: #2e7d32;
        color: white;
        border-radius: 8px;
    }}
    .stButton button {{
        background-color: #388e3c;
        color: white;
        border-radius: 8px;
    }}
    .logo-container {{
        position: fixed;
        top: 0px;
        left: 0px;
        z-index: 999;
    }}
    .logo-container img {{
        width: 80px;
        opacity: 0.95;
    }}
    </style>
    <div class="logo-container">
        <img src="data:image/png;base64,{logo_base64}" alt="شعار متوازي">
    </div>
    """,
    unsafe_allow_html=True
)

# عرض العنوان والنموذج
st.title(TEXTS[lang]["title"])
st.markdown(TEXTS[lang]["instruction"])

uploaded_file = st.file_uploader(TEXTS[lang]["upload_label"], type=["pdf"])
project_name = st.text_input(TEXTS[lang]["project_name"])
client_name = st.text_input(TEXTS[lang]["client_name"])
gov_logo = st.file_uploader(TEXTS[lang]["gov_logo"], type=["png", "jpg"])

@st.cache_data
def extract_text_from_pdf(file):
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def generate_proposal(content, project, client_name, lang):
    if lang == "العربية":
        system_prompt = "أنت مساعد خبير في كتابة العروض الفنية بناءً على كراسة الشروط."
        user_prompt = f"""هذه كراسة شروط لمشروع جديد: {content}

اكتب عرضًا فنيًا متكاملًا باسم المشروع: {project} والجهة: {client_name}، وابدأه بمقدمة من نحن، ثم فهم المشروع، نطاق العمل، المنهجية، الجدول الزمني، الفريق المقترح، والخاتمة.
"""
    else:
        system_prompt = "You are an expert assistant for writing professional technical proposals based on an RFP."
        user_prompt = f"""This is an RFP document: {content}

Write a full technical proposal for the project titled: {project}, and client: {client_name}. Start with a company introduction, then project understanding, scope of work, methodology, timeline, proposed team, and conclusion.
"""

    response = client.chat.completions.create(
        model="gpt-4-turbo-2024-04-09",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        max_tokens=3000,
        temperature=0.7
    )
    return response.choices[0].message.content

if st.button(TEXTS[lang]["generate_button"]):
    if uploaded_file and project_name and client_name:
        with st.spinner(TEXTS[lang]["loading"]):
            extracted_text = extract_text_from_pdf(uploaded_file)
            proposal_text = generate_proposal(extracted_text, project_name, client_name, lang)

        doc = Document()
        doc.add_heading(f"{TEXTS[lang]['doc_heading']} - {project_name}", level=1)
        doc.add_paragraph(f"{TEXTS[lang]['client_name']}: {client_name}")

        for paragraph in proposal_text.split("\n"):
            line = paragraph.strip()
            if line.startswith("### "):
                doc.add_heading(line.replace("### ", ""), level=3)
            elif line.startswith("## "):
                doc.add_heading(line.replace("## ", ""), level=2)
            elif line.startswith("# "):
                doc.add_heading(line.replace("# ", ""), level=1)
            elif line:
                doc.add_paragraph(line)

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        with open(tmp_path, "rb") as f:
            st.download_button(TEXTS[lang]["download"], f, file_name=f"Proposal_{project_name}.docx")

        st.success(TEXTS[lang]["success"])
    else:
        st.error(TEXTS[lang]["error"])
